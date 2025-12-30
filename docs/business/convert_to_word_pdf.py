#!/usr/bin/env python3
"""
Script to convert AWS Credits Pitch Deck Markdown to Word and PDF formats.

Requirements:
    pip install pypandoc python-docx markdown pdfkit

For PDF generation, you also need wkhtmltopdf:
    macOS: brew install wkhtmltopdf
    Linux: sudo apt-get install wkhtmltopdf
    Windows: Download from https://wkhtmltopdf.org/downloads.html
"""

import os
import sys
from pathlib import Path

def convert_with_pypandoc():
    """Convert using pypandoc (recommended method)."""
    try:
        import pypandoc
        
        md_file = Path(__file__).parent / "AWS_CREDITS_PITCH_DECK.md"
        docx_file = md_file.with_suffix('.docx')
        pdf_file = md_file.with_suffix('.pdf')
        
        print(f"Converting {md_file} to Word format...")
        pypandoc.convert_file(
            str(md_file),
            'docx',
            outputfile=str(docx_file),
            extra_args=['--reference-doc=reference.docx'] if Path('reference.docx').exists() else []
        )
        print(f"✓ Created: {docx_file}")
        
        print(f"Converting {md_file} to PDF format...")
        pypandoc.convert_file(
            str(md_file),
            'pdf',
            outputfile=str(pdf_file),
            extra_args=['--pdf-engine=wkhtmltopdf', '--toc']
        )
        print(f"✓ Created: {pdf_file}")
        
        return True
    except ImportError:
        print("Error: pypandoc not installed. Install with: pip install pypandoc")
        return False
    except Exception as e:
        print(f"Error converting with pypandoc: {e}")
        return False

def convert_with_markdown_docx():
    """Convert using markdown and python-docx (alternative method)."""
    try:
        import markdown
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        md_file = Path(__file__).parent / "AWS_CREDITS_PITCH_DECK.md"
        docx_file = md_file.with_suffix('.docx')
        
        print(f"Reading {md_file}...")
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to HTML
        html = markdown.markdown(md_content, extensions=['extra', 'toc'])
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Plixera - AWS Credits Application', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_paragraph('DevBridge: Mobile App Monitoring & Configuration Platform')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Parse and add content (simplified - for better results use pypandoc)
        lines = md_content.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.strip():
                doc.add_paragraph(line)
        
        doc.save(str(docx_file))
        print(f"✓ Created: {docx_file}")
        print("Note: This is a basic conversion. For better formatting, use pypandoc.")
        
        return True
    except ImportError as e:
        print(f"Error: Required library not installed: {e}")
        print("Install with: pip install markdown python-docx")
        return False
    except Exception as e:
        print(f"Error converting: {e}")
        return False

def main():
    """Main conversion function."""
    print("=" * 60)
    print("AWS Credits Pitch Deck Converter")
    print("=" * 60)
    print()
    
    # Try pypandoc first (best quality)
    if convert_with_pypandoc():
        print("\n✓ Conversion completed successfully!")
        return
    
    # Fallback to markdown+docx
    print("\nTrying alternative conversion method...")
    if convert_with_markdown_docx():
        print("\n✓ Basic conversion completed!")
        print("For better formatting, install pypandoc: pip install pypandoc")
        return
    
    print("\n✗ Conversion failed. Please install required dependencies:")
    print("  pip install pypandoc python-docx markdown pdfkit")
    print("\nOr use online converters:")
    print("  - https://www.markdowntopdf.com/")
    print("  - https://www.zamzar.com/convert/md-to-docx/")
    print("  - https://cloudconvert.com/md-to-docx")

if __name__ == '__main__':
    main()




