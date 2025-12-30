#!/usr/bin/env python3
"""
Generate Word document from Markdown pitch deck.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from pathlib import Path

def create_word_document():
    """Create Word document from markdown file."""
    
    # Read markdown file
    md_file = Path(__file__).parent / "AWS_CREDITS_PITCH_DECK.md"
    docx_file = md_file.with_suffix('.docx')
    
    print(f"Reading {md_file}...")
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Parse markdown and add to document
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines (but add spacing)
        if not line:
            i += 1
            continue
        
        # Title (first line)
        if i == 0 and line.startswith('# '):
            title = doc.add_heading(line[2:], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_format = title.runs[0].font
            title_format.size = Pt(24)
            title_format.bold = True
            title_format.color.rgb = RGBColor(26, 115, 232)  # Blue
            i += 1
            continue
        
        # Main headings (# Heading)
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], 1)
            heading_format = heading.runs[0].font
            heading_format.color.rgb = RGBColor(26, 115, 232)  # Blue
            i += 1
            continue
        
        # Subheadings (## Subheading)
        if line.startswith('## '):
            heading = doc.add_heading(line[3:], 2)
            heading_format = heading.runs[0].font
            heading_format.color.rgb = RGBColor(66, 133, 244)  # Light blue
            i += 1
            continue
        
        # Sub-subheadings (### Sub-subheading)
        if line.startswith('### '):
            heading = doc.add_heading(line[4:], 3)
            heading_format = heading.runs[0].font
            heading_format.color.rgb = RGBColor(95, 99, 104)  # Gray
            i += 1
            continue
        
        # Bold text (**text**)
        if line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            i += 1
            continue
        
        # Lists (- item or * item)
        if line.startswith('- ') or line.startswith('* '):
            items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                item_text = lines[i].strip()[2:].strip()
                # Remove markdown formatting
                item_text = re.sub(r'\*\*(.*?)\*\*', r'\1', item_text)
                items.append(item_text)
                i += 1
            
            for item in items:
                p = doc.add_paragraph(item, style='List Bullet')
            continue
        
        # Numbered lists (1. item)
        if re.match(r'^\d+\.\s', line):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                item_text = re.sub(r'^\d+\.\s', '', lines[i].strip())
                item_text = re.sub(r'\*\*(.*?)\*\*', r'\1', item_text)
                items.append(item_text)
                i += 1
            
            for item in items:
                p = doc.add_paragraph(item, style='List Number')
            continue
        
        # Tables (basic support)
        if '|' in line and line.count('|') >= 2:
            # Skip separator lines
            if re.match(r'^\|[\s\-:]+\|', line):
                i += 1
                continue
            
            # Parse table
            rows = []
            while i < len(lines) and '|' in lines[i]:
                if not re.match(r'^\|[\s\-:]+\|', lines[i].strip()):
                    cells = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                    rows.append(cells)
                i += 1
            
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        # Remove markdown formatting
                        cell_data = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_data)
                        table.rows[row_idx].cells[col_idx].text = cell_data
                        # Make header row bold
                        if row_idx == 0:
                            for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
            continue
        
        # Regular paragraph
        # Remove markdown formatting
        para_text = line
        para_text = re.sub(r'\*\*(.*?)\*\*', r'\1', para_text)  # Bold
        para_text = re.sub(r'`(.*?)`', r'\1', para_text)  # Code
        
        # Check if it's a special line (like "Date:", "Company:", etc.)
        if ':' in para_text and len(para_text.split(':')) == 2:
            p = doc.add_paragraph()
            parts = para_text.split(':', 1)
            run1 = p.add_run(parts[0] + ':')
            run1.bold = True
            p.add_run(parts[1])
        else:
            p = doc.add_paragraph(para_text)
        
        i += 1
    
    # Save document
    print(f"Saving Word document to {docx_file}...")
    doc.save(str(docx_file))
    print(f"✓ Successfully created: {docx_file}")
    return docx_file

if __name__ == '__main__':
    try:
        create_word_document()
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()




